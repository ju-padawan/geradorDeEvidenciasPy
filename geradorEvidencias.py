import docx
from docx.text.run import *
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import yaml
import os


with open('data/dados.yaml') as arquivo:
    dados = yaml.load(arquivo, Loader=yaml.FullLoader)

#Deescobrindo quantos cenários foram executados
lista_cen = len(dados)

#Percorrer a lista de cenários
count = 1
while count <= lista_cen:
    #Descobrindo quantos campos tem em cada cenário
    cen = 'cenario_'+str(count)
    lista_dados = len(dados[cen])
    pasta_evidencias = dados[cen]['pasta_evidencias']

    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    doc.add_heading(str(dados[cen]['cen_nome']), 0)

    p = doc.add_paragraph()
    p.add_run('Descrição: ').bold = True
    p.add_run(str(dados[cen]['cen_desc']))

    p = doc.add_paragraph()
    p.add_run('Pré-requisitos: ').bold = True
    p.add_run(str(dados[cen]['cen_pre_requisitos']))

    p = doc.add_paragraph()
    p.add_run('Resultado esperado: ').bold = True
    p.add_run(str(dados[cen]['resultado_esperado']))

    p = doc.add_paragraph()
    p.add_run('Plataforma: ').bold = True
    p.add_run(str(dados[cen]['cen_plataforma']))

    p = doc.add_paragraph()
    p.add_run('Disposiitivo uilizado: ').bold = True
    p.add_run(str(dados[cen]['cen_dispositivo']))

    p = doc.add_paragraph()
    p.add_run('Versão do Software: ').bold = True
    p.add_run(str(dados[cen]['cen_versao_software']))

    p = doc.add_paragraph()
    p.add_run('Versão do App Next: ').bold = True
    p.add_run(str(dados[cen]['cen_versao_app']))

    p = doc.add_paragraph()
    p.add_run('Executado por: ').bold = True
    p.add_run(str(dados[cen]['cen_executor']))

    p = doc.add_paragraph()
    p.add_run('Massa uilizada: ').bold = True
    p.add_run(str(dados[cen]['cen_massa']))

    p = doc.add_paragraph()
    p.add_run('Data da execução: ').bold = True
    p.add_run(str(dados[cen]['cen_data_execucao']))

    p = doc.add_paragraph()
    p.add_run('Status execução: ').bold = True
    p.add_run(str(dados[cen]['cen_status_execucao']))
    #p = doc.add_paragraph().paragraph_format.space_before = Pt(12)
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)

    #percorrer os campos de cada cenário
    passos = (lista_dados - 13) #/2
    count1 = 1
    while count1 <= int(passos):
        cen_passo = 'passo_'+str(count1)
        #cen_resultado = 'resultado_'+str(count1)

        evidencia = pasta_evidencias+'/'+cen_passo+'.png'

        txt_passo = '[Passo_'+str(count1)+']: '

        p = doc.add_paragraph()
        p.add_run(txt_passo).bold = True
        p.add_run(str(dados[cen][cen_passo]))
        p1 = doc.add_paragraph()
        p1.add_run("Resultado: ").bold = True
        p2 = doc.add_paragraph().paragraph_format.space_before = Pt(4)
        
        #p.add_run(str(dados[cen][cen_resultado]))

        doc.add_picture(evidencia, width=Inches(6))

        if count1 < int(passos):
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_break(WD_BREAK.PAGE)
    
        #os.remove(evidencia)

        count1 = count1+1

    nome_cenario = str(dados[cen]['cen_nome'])
    path_doc = str(dados[cen]['pasta_evidencias'])+'/'+nome_cenario[0:21]+' - '+str(dados[cen]['cen_plataforma'])+'.docx'
    doc.save(path_doc)
    print("==============================================================\n")
    print("Evidência "+nome_cenario[0:21]+" Gerada com sucesso!")
    print("\n==============================================================")


    count = count+1

