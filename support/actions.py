import docx
from docx.text.run import *
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor
import yaml
import os

class SupportActions():

    def __init__(self) -> None:
        pass

    def ler_dados_arquivo_yaml(self):
        try:
            with open('data/dados.yaml') as arquivo:
                dados = yaml.load(arquivo, Loader=yaml.FullLoader)
                return dados
        except BaseException as error:
            print(f"Unexpected {error=}, {type(error)=}")
            raise

    def formatacao_arquivo_evidencia(self):
        self.doc = Document()
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(12)

    def inserir_titulo_arquivo_evidencia(self, titulo):
        self.doc.add_heading(titulo, 0)

    def inserir_cabecalho_arquivo_evidencia(self, posicao,  cabecalho, tipo_cabecalho):
        section = self.doc.sections[posicao]
        header = section.header
        header_para = header.paragraphs[posicao]
        if tipo_cabecalho == "centralizado":
            header_para.text = "\t"+cabecalho
        elif tipo_cabecalho == "direita":
            header_para.text = "\t\t"+cabecalho
        else:
            header_para.text = cabecalho

    def inserir_imagem_cabecalho_arquivo_evidencia(self, posicao,  cabecalho_imagem):
        section = self.doc.sections[posicao]
        header = section.header
        header_para = header.paragraphs[posicao]
        r = header_para.add_run()
        r.add_picture(cabecalho_imagem)

    def inserir_rodape_arquivo_evidencia(self, posicao,  rodape, tipo_rodape):
        section = self.doc.sections[posicao]
        footer = section.footer
        footer_para = footer.paragraphs[posicao]
        if tipo_rodape == "centralizado":
            footer_para.text = "\t"+rodape
        elif tipo_rodape == "direita":
            footer_para.text = "\t\t"+rodape
        else:
            footer_para.text =rodape
        

    def inserir_informacoes_arquivo_evidencia(self, texto, dado):
        paragrafo = self.doc.add_paragraph()
        paragrafo.add_run(texto).bold = True
        paragrafo.add_run(dado)

    def inserir_quebra_de_pagina(self):
        p = self.doc.add_paragraph()
        run = p.add_run()
        run.add_break(WD_BREAK.PAGE)

    def inserir_espaco_antes_paragrafo(self, espaco):
        self.doc.add_paragraph().paragraph_format.space_after = Pt(espaco)

    def inserir_espaco_apos_paragrafo(self, espaco):
        self.doc.add_paragraph().paragraph_format.space_before = Pt(espaco)


    def inserir_imagem_arquivo_evidencia(self, imagem, msg_erro):
        if os.path.exists(imagem):
            self.doc.add_picture(imagem, width=Inches(5.75))
        else:
            print(msg_erro)


    def salvar_arquivo_evidencia(self, nome_cenario, path, status, plataforma):
        try:
            if "Sucesso" in status:
                path_doc = path+'/'+nome_cenario+"_"+plataforma+"_Sucesso"+".docx"
            else:
                path_doc = path+'/'+nome_cenario+"_"+plataforma+"_Falha"+".docx"
                
            self.doc.save(path_doc)
            print("\n>>>> Evidência "+nome_cenario+" Gerada com sucesso! <<<<")
            print("\n==================================================================================")
        except BaseException as error:
            print(f"Unexpected {error=}, {type(error)=}")
            raise

    def exibir_informacao_console(self, texto):
        print(texto)

    def inserir_status_colorido(self, texto, status):
        if "Sucesso" in status:
            paragrafo = self.doc.add_paragraph()
            paragrafo.add_run(texto).bold = True
            run = paragrafo.add_run(status)
            font = run.font
            font.color.rgb = RGBColor(0, 128, 0)

        elif "Falha" in status:
            paragrafo = self.doc.add_paragraph()
            paragrafo.add_run(texto).bold = True
            run = paragrafo.add_run(status)
            font = run.font
            font.color.rgb = RGBColor(255, 0, 0)

    

    